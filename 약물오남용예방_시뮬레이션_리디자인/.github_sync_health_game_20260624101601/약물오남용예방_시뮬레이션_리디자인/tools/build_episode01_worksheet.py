from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "worksheets" / "에피소드1_떨리는손의이유_학생활동지_v1.docx"

PAGE_W = 11906
PAGE_H = 16838
MARGIN_LR = 760
MARGIN_TOP = 700
MARGIN_BOTTOM = 560
CONTENT_W = 10386

FONT = "맑은 고딕"
MAIN = "D7F0FF"
ACCENT = "B6E2FB"
WARNING = "FFE1EC"
WRITE = "F9FAFB"
BORDER = "D1D5DB"
TEXT = "1A1A1A"
BLUE = "1A5C8A"
LABEL_BLUE = "2E86C1"


def set_rfonts(run, font=FONT):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font)


def set_run(run, size=10.5, bold=False, color=TEXT):
    set_rfonts(run)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_p(p, after=4, before=0, line=1.15, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def add_para(doc, text="", size=10.5, bold=False, color=TEXT, after=4, before=0, align=None):
    p = doc.add_paragraph()
    set_p(p, after=after, before=before, align=align)
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, color=color)
    return p


def set_style_font(style, size=10.5, bold=False, color=TEXT):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), FONT)


def set_table_geometry(table, widths, border_color=BORDER, fill_header=None):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblBorders", "w:tblCellMar"):
        for old in tbl_pr.findall(qn(tag)):
            tbl_pr.remove(old)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), border_color)
        borders.append(el)
    tbl_pr.append(borders)

    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", 70), ("bottom", 70), ("left", 120), ("right", 120)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    tbl.insert(1, grid)

    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            if fill_header and row_idx == 0:
                shade_cell(cell, fill_header)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def format_cell(cell, bold=False, color=TEXT, size=10.0):
    for p in cell.paragraphs:
        set_p(p, after=2, line=1.15)
        for run in p.runs:
            set_run(run, size=size, bold=bold, color=color)


def clear_cell(cell):
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    return p


def add_box(doc, body, fill=MAIN, border="C9DEEC", label=None):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_W], border_color=border)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = clear_cell(cell)
    set_p(p, after=2, line=1.15)
    if label:
        r = p.add_run(label)
        set_run(r, bold=True, color=BLUE)
        p.add_run("\n")
    for idx, line in enumerate(body.split("\n")):
        r = p.add_run(line)
        set_run(r, size=10.5, color=TEXT)
        if idx != len(body.split("\n")) - 1:
            p.add_run("\n")
    add_para(doc, "", after=2)
    return table


def add_write_box(doc, label, prompt_lines, blanks=2):
    body = "\n".join(prompt_lines + ["_" * 62 for _ in range(blanks)])
    return add_box(doc, body, fill=WRITE, border=BORDER, label=f"✎ {label}")


def add_heading(doc, text, level=2):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(14 if level == 2 else 0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, size=12 if level == 2 else 17, bold=True, color=BLUE if level == 2 else TEXT)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_geometry(table, widths, border_color=BORDER, fill_header=ACCENT)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            format_cell(cell, bold=(i == 0), size=9.6 if len(headers) >= 4 else 10.0)
    add_para(doc, "", after=2)
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(PAGE_W)
    sec.page_height = Twips(PAGE_H)
    sec.top_margin = Twips(MARGIN_TOP)
    sec.bottom_margin = Twips(MARGIN_BOTTOM)
    sec.left_margin = Twips(MARGIN_LR)
    sec.right_margin = Twips(MARGIN_LR)
    sec.header_distance = Twips(708)
    sec.footer_distance = Twips(708)

    set_style_font(doc.styles["Normal"], 10.5, False, TEXT)
    set_style_font(doc.styles["Heading 1"], 17, True, TEXT)
    set_style_font(doc.styles["Heading 2"], 12, True, BLUE)

    add_para(doc, "에피소드 1 · 시험기간 카페인 사건", size=9, bold=True, color=LABEL_BLUE, after=2)
    add_heading(doc, "떨리는 손의 이유 — 건강 미스터리 사건 기록지", level=1)
    add_para(doc, "학번 __________   이름 __________", align=WD_ALIGN_PARAGRAPH.RIGHT, after=8)

    add_box(
        doc,
        "보건 미스터리 파일 첫 사건입니다. 민재의 손 떨림과 창백한 얼굴을 단순한 피곤함으로 넘기지 말고, "
        "관찰한 사실과 단서를 연결해 원인을 조합해 봅니다. 정답 하나를 맞히기보다 근거를 들어 설명하는 것이 목표입니다.",
        fill=MAIN,
        border="C9DEEC",
        label="🔎 수사반장 브리핑",
    )

    add_heading(doc, "준비 1. 사건 브리핑 — 몸이 보내는 신호 읽기")
    add_box(
        doc,
        "시험을 앞둔 오후, 민재는 괜찮다고 말하지만 손을 계속 떨고 얼굴도 하얘 보입니다. "
        "건강 신호는 작은 변화처럼 보여도 여러 원인이 겹쳤다는 단서가 될 수 있습니다.",
        fill=MAIN,
        border="C9DEEC",
    )
    add_write_box(
        doc,
        "수사 기록란",
        [
            "민재에게 보인 이상 징후를 2가지 이상 적어보세요.",
            "처음 떠오른 원인을 적어보세요. 아직 확인되지 않은 추측이어도 괜찮습니다.",
        ],
        blanks=3,
    )

    add_heading(doc, "준비 2. 사실과 추측 구분하기")
    add_box(
        doc,
        "직접 보거나 확인한 것은 '사실', 아직 가능성으로 생각하는 것은 '추측'입니다. "
        "성급히 한 가지 원인으로 단정하지 않는 것이 건강 추리의 첫 원칙입니다.",
        fill=ACCENT,
        border="9FD1ED",
        label="📌 기억할 점",
    )
    add_table(
        doc,
        ["관찰한 사실", "아직 추측인 점"],
        [
            ["예: 손을 떨고 있었다", "예: 잠을 못 자서 그런 것일 수 있다"],
            ["", ""],
            ["", ""],
        ],
        [5193, 5193],
    )

    add_heading(doc, "1단계. 단서 수집하기")
    add_box(
        doc,
        "게임 속 조사 장면에서 찾은 단서를 아래 표에 정리하세요. 단서 이름만 쓰지 말고, "
        "그 단서가 왜 중요한지도 함께 적어야 합니다.",
        fill=MAIN,
        border="C9DEEC",
    )
    add_table(
        doc,
        ["단서", "확인한 내용", "왜 중요한가?"],
        [
            ["감기약 봉투", "아침 복용 표시가 있다", ""],
            ["메모지", "'아침 약', '졸리면 편의점' 같은 메모가 있다", ""],
            ["카페인 음료 캔", "카페인 음료를 마신 흔적이 있다", ""],
            ["편의점 영수증", "오전과 점심 시간대 구매 기록이 있다", ""],
            ["메신저 캡처", "'안 마시면 못 버틸 듯'이라는 말이 보인다", ""],
        ],
        [2100, 4300, 3986],
    )
    add_box(
        doc,
        "한 가지 단서만으로 결론을 내리면 실제 원인을 놓칠 수 있습니다. 약 복용, 수면 부족, 카페인 섭취는 따로 보지 말고 함께 봐야 합니다.",
        fill=WARNING,
        border="F4B8CF",
        label="⚠ 함정 주의",
    )

    add_heading(doc, "2단계. 가설 세우고 수정하기")
    add_write_box(
        doc,
        "처음 가설",
        ["처음에는 민재의 상태가 왜 나빠졌다고 생각했나요?"],
        blanks=2,
    )
    add_write_box(
        doc,
        "생각이 바뀐 이유",
        ["새로 확인한 단서 때문에 생각이 어떻게 바뀌었나요? 근거가 된 단서 이름을 포함하세요."],
        blanks=3,
    )
    add_table(
        doc,
        ["원인 후보", "확인 여부", "근거"],
        [
            ["수면 부족", "□ 관련 있음  □ 확실하지 않음", ""],
            ["감기약 복용", "□ 관련 있음  □ 확실하지 않음", ""],
            ["카페인 반복 섭취", "□ 관련 있음  □ 확실하지 않음", ""],
            ["시험 압박과 버티려는 마음", "□ 관련 있음  □ 확실하지 않음", ""],
        ],
        [2600, 3000, 4786],
    )

    add_heading(doc, "3단계. 안전한 대응 정하기")
    add_box(
        doc,
        "친구를 혼내거나 단정하기보다, 무엇을 언제 얼마나 먹었는지 확인하고 도움을 요청하는 것이 안전합니다. "
        "몸의 이상 신호가 계속되면 보건실, 보호자, 전문가에게 연결해야 합니다.",
        fill=ACCENT,
        border="9FD1ED",
        label="📌 안전 대응 원칙",
    )
    add_table(
        doc,
        ["상황", "내가 할 질문", "안전한 대응"],
        [
            ["친구가 괜찮다고 하지만 얼굴이 창백하다", "오늘 먹은 약이나 음료가 있었는지 묻기", ""],
            ["약 복용과 카페인 음료가 함께 확인된다", "언제, 얼마나, 몇 번 먹었는지 확인하기", ""],
            ["손 떨림·불안감이 계속된다", "혼자 버티지 않도록 보건실 도움 요청하기", ""],
        ],
        [2900, 3900, 3586],
    )
    add_write_box(
        doc,
        "내가 선택할 대응",
        ["민재에게 실제로 어떤 말을 하고, 누구에게 도움을 요청할지 적어보세요."],
        blanks=2,
    )

    add_heading(doc, "4단계. 근거로 설명하고 역량 확인하기")
    add_write_box(
        doc,
        "최종 원인 조합",
        ["민재의 상태를 만든 원인을 한 가지가 아니라 조합으로 설명해보세요."],
        blanks=3,
    )
    add_table(
        doc,
        ["내가 사용한 역량", "체크", "이번 사건에서 한 행동"],
        [
            ["관찰력", "□", "증상과 행동 변화를 보았다"],
            ["분석력", "□", "단서를 연결해 원인 조합을 생각했다"],
            ["공감력", "□", "민재가 버티려 했던 이유를 이해했다"],
            ["신중함", "□", "한 가지 원인으로 단정하지 않았다"],
            ["대응력", "□", "보건실 도움 요청을 선택했다"],
            ["건강문해력", "□", "약 복용과 카페인 섭취를 함께 고려했다"],
        ],
        [2500, 900, 6986],
    )
    add_write_box(
        doc,
        "오늘 배운 점",
        ["내 생활에서 비슷한 상황이 생기면 무엇을 먼저 확인해야 할까요?"],
        blanks=2,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
